"""DOCX text extraction using python-docx with embedded image OCR."""

import io
import os
from models import ExtractionResult, PageContent, ExtractionMetadata
from .base import BaseExtractor
from utils.text_normalizer import normalize_text

try:
    import pytesseract
    TESSERACT_AVAILABLE = True
except ImportError:
    TESSERACT_AVAILABLE = False


class DOCXExtractor(BaseExtractor):
    """Extract text from DOCX files, including OCR on embedded images."""

    async def extract(self, file_path: str, file_name: str) -> ExtractionResult:
        try:
            from docx import Document
            from docx.opc.constants import RELATIONSHIP_TYPE as RT
            from PIL import Image
        except ImportError as e:
            return self._create_error_result(file_name, "docx", f"Missing dependency: {e}")

        file_size = os.path.getsize(file_path)

        try:
            doc = Document(file_path)

            # ---- Extract paragraph text ----
            paragraphs: list[str] = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)

            # ---- Extract table text ----
            has_tables = len(doc.tables) > 0
            table_texts: list[str] = []
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    table_texts.append(" | ".join(cells))

            # ---- Extract images and run OCR ----
            image_texts: list[str] = []
            has_images = False
            if TESSERACT_AVAILABLE:
                for rel in doc.part.rels.values():
                    if "image" in rel.reltype:
                        has_images = True
                        try:
                            image_data = rel.target_part.blob
                            img = Image.open(io.BytesIO(image_data))
                            text = pytesseract.image_to_string(img)
                            if text.strip():
                                image_texts.append(text.strip())
                        except Exception:
                            continue

            # ---- Build combined text ----
            all_parts: list[str] = []
            all_parts.extend(paragraphs)
            if table_texts:
                all_parts.append("\n[TABLES]\n" + "\n".join(table_texts))
            if image_texts:
                all_parts.append("\n[IMAGE TEXT]\n" + "\n".join(image_texts))

            extracted = normalize_text("\n\n".join(all_parts))

            # Treat entire document as single page
            page = PageContent(
                page_number=1,
                text=extracted,
                has_images=has_images,
                image_text="\n".join(image_texts) if image_texts else "",
            )

            word_count = len(extracted.split())

            return ExtractionResult(
                success=True,
                file_name=file_name,
                file_type="docx",
                file_size=file_size,
                total_pages=1,
                extracted_text=extracted,
                pages=[page],
                metadata=ExtractionMetadata(
                    word_count=word_count,
                    char_count=len(extracted),
                    has_images=has_images,
                    has_tables=has_tables,
                    extraction_method="python-docx+OCR",
                ),
            )
        except Exception as e:
            return self._create_error_result(file_name, "docx", str(e))
